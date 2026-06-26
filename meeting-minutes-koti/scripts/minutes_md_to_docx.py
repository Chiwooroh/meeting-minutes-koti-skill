#!/usr/bin/env python
"""Fill the Korean meeting-minutes DOCX template from a minutes Markdown file."""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


FIELD_NAMES = ["과제명", "회의명", "개최일시", "장소", "참석인원", "회의목적"]
FIELD_ALIASES = {
    "과제명": {"과제명", "과 제 명", "怨쇱젣紐?"},
    "회의명": {"회의명", "회 의 명", "?뚯쓽紐?"},
    "개최일시": {"개최일시", "일시", "媛쒖턀?쇱떆"},
    "장소": {"장소", "?μ냼"},
    "참석인원": {"참석인원", "참 석 인 원", "李몄꽍?몄썝"},
    "회의목적": {"회의목적", "회 의 목 적", "?뚯쓽紐⑹쟻"},
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create DOCX meeting minutes from Markdown.")
    parser.add_argument("--template", required=True, help="DOCX template path.")
    parser.add_argument("--input", required=True, help="Minutes Markdown path.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    return parser.parse_args()


def normalize_label(text: str) -> str:
    return re.sub(r"\s+", "", text.strip())


def canonical_field(label: str) -> str | None:
    compact = normalize_label(label)
    for field, aliases in FIELD_ALIASES.items():
        if compact in {normalize_label(alias) for alias in aliases}:
            return field
    return None


def parse_minutes(markdown: str) -> tuple[dict[str, str], str]:
    fields = {name: "" for name in FIELD_NAMES}

    for line in markdown.splitlines():
        if not line.startswith("|"):
            continue
        cols = [part.strip() for part in line.strip().strip("|").split("|")]
        if len(cols) < 2 or re.fullmatch(r"[-:\s]+", cols[0]) or cols[0] in {"구분", "援щ텇"}:
            continue
        field = canonical_field(cols[0])
        if field:
            fields[field] = cols[1]

    content_patterns = [
        r"##\s*회의\s*내용\s*(.*?)(?=\n##\s*공지|\n##\s*확인|\Z)",
        r"##\s*\?뚯쓽\?댁슜\s*(.*?)(?=\n##\s*怨듭|\n##\s*\?뺤씤|\Z)",
    ]
    content = ""
    for pattern in content_patterns:
        match = re.search(pattern, markdown, re.S)
        if match:
            content = match.group(1).strip()
            break

    check_patterns = [
        r"##\s*확인\s*필요\s*사항\s*(.*?)(?=\n## |\Z)",
        r"##\s*\?뺤씤\s*\?꾩슂\s*\?ы빆\s*(.*?)(?=\n## |\Z)",
    ]
    for pattern in check_patterns:
        match = re.search(pattern, markdown, re.S)
        if not match:
            continue
        check_text = match.group(1).strip()
        if check_text and check_text not in {"- 없음", "- ?놁쓬"}:
            content = f"{content}\n\n□ 확인 필요 사항\n{check_text}".strip()
        break

    return fields, content


def clear_cell(cell) -> None:
    cell.text = ""


def write_cell(cell, text: str, font_size: int = 10) -> None:
    clear_cell(cell)
    lines = text.splitlines() or [""]
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(line)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(font_size)


def fill_template(template: Path, output: Path, fields: dict[str, str], content: str) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template, output)
    doc = Document(str(output))
    if not doc.tables:
        raise SystemExit("Template has no table to fill.")
    table = doc.tables[0]
    if len(table.rows) < 6:
        raise SystemExit("Template table must have at least 6 rows.")

    write_cell(table.rows[0].cells[1], fields["과제명"])
    write_cell(table.rows[1].cells[1], fields["회의명"])
    write_cell(table.rows[2].cells[2], fields["개최일시"])
    write_cell(table.rows[2].cells[4], fields["장소"])
    write_cell(table.rows[3].cells[1], fields["참석인원"])
    write_cell(table.rows[4].cells[1], fields["회의목적"])
    write_cell(table.rows[5].cells[1], content, font_size=9)
    doc.save(str(output))


def main() -> None:
    args = parse_args()
    template = Path(args.template).expanduser().resolve()
    input_path = Path(args.input).expanduser().resolve()
    output = Path(args.output).expanduser().resolve()
    fields, content = parse_minutes(input_path.read_text(encoding="utf-8"))
    fill_template(template, output, fields, content)
    print(f"Wrote DOCX: {output}")


if __name__ == "__main__":
    main()
